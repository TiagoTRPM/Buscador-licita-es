# etp_gerador.py
import io
import os
from datetime import datetime
from pathlib import Path


def gerar_word(etp: dict) -> bytes:
    """
    Gera o ETP em formato Word (.docx) seguindo o padrão da Lei 14.133/2021.
    Retorna os bytes do arquivo.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("Instale python-docx: pip install python-docx")

    doc = Document()

    # ── Configuração da página ──────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    cab  = etp.get("cabecalho", {})
    fixo = etp.get("campos_fixos", {})
    itens = etp.get("itens", [])
    levantamento = etp.get("levantamento", [])
    campos_livres = etp.get("campos_livres", [])

    orgao     = cab.get("orgao", "")
    municipio = cab.get("municipio", "")
    estado    = cab.get("estado", "")
    processo  = cab.get("processo", "")
    responsavel = cab.get("responsavel", "")
    cargo     = cab.get("cargo", "")
    data_str  = cab.get("data", datetime.now().strftime("%Y-%m-%d"))
    titulo_etp = etp.get("titulo", "Estudo Técnico Preliminar")

    try:
        dt = datetime.strptime(data_str, "%Y-%m-%d")
        data_fmt = dt.strftime("%d de %B de %Y").replace(
            "January","janeiro").replace("February","fevereiro").replace(
            "March","março").replace("April","abril").replace(
            "May","maio").replace("June","junho").replace(
            "July","julho").replace("August","agosto").replace(
            "September","setembro").replace("October","outubro").replace(
            "November","novembro").replace("December","dezembro")
    except Exception:
        data_fmt = data_str

    # ── Helpers ─────────────────────────────────────

    def set_font(run, size=12, bold=False, color=None):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def titulo_secao(texto, numero=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(f"{numero} {texto}" if numero else texto)
        set_font(run, size=12, bold=True)
        return p

    def corpo(texto, negrito=False, italico=False, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY):
        if not texto:
            texto = "Não informado."
        p = doc.add_paragraph()
        p.alignment = alinhamento
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(texto)
        set_font(run, size=12)
        run.font.bold   = negrito
        run.font.italic = italico
        return p

    def linha_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Cabeçalho ───────────────────────────────────
    modo = cab.get("modo_cabecalho", "dinamico")
    brasao_path = cab.get("brasao_path", "")

    if modo == "arquivo" and cab.get("cabecalho_path") and os.path.exists(cab["cabecalho_path"]):
        from docx.shared import Inches
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(cab["cabecalho_path"], width=Inches(6))
    else:
        # Cabeçalho dinâmico
        if brasao_path and os.path.exists(brasao_path):
            from docx.shared import Inches
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(brasao_path, height=Cm(2.5))

        if orgao:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(orgao.upper())
            set_font(run, size=14, bold=True)

        if municipio or estado:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{municipio} — {estado}" if estado else municipio)
            set_font(run, size=11)

    linha_hr()
    doc.add_paragraph()

    # ── Título principal ─────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESTUDO TÉCNICO PRELIMINAR")
    set_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo_etp)
    set_font(run, size=13, bold=True)

    if processo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Processo nº {processo}")
        set_font(run, size=11)

    doc.add_paragraph()
    linha_hr()

    # ── Introdução ───────────────────────────────────
    titulo_secao("INTRODUÇÃO")
    corpo(
        "O presente documento apresenta a etapa de planejamento e os devidos estudos para a "
        "contratação de solução que atenderá à necessidade abaixo especificada, em conformidade "
        "com a Lei Federal nº 14.133/2021 e suas regulamentações."
    )

    # ── 1. Descrição da Necessidade ──────────────────
    titulo_secao("DESCRIÇÃO DA NECESSIDADE", "1 —")
    corpo(fixo.get("descricao_necessidade", ""))

    # ── 2. Previsão no PCA ───────────────────────────
    titulo_secao("PREVISÃO NO PLANO DE CONTRATAÇÕES ANUAL", "2 —")
    corpo(fixo.get("previsao_pca", ""))

    # ── 3. Requisitos da Contratação ─────────────────
    titulo_secao("REQUISITOS DA CONTRATAÇÃO", "3 —")
    corpo(fixo.get("requisitos_contratacao", ""))

    # ── 4. Estimativa das Quantidades ────────────────
    titulo_secao("ESTIMATIVA DAS QUANTIDADES", "4 —")

    if itens:
        tabela = doc.add_table(rows=1, cols=4)
        tabela.style = "Table Grid"
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Larguras em DXA (1440 = 1 polegada)
        larguras = [720, 2160, 5040, 1080]  # Nº, Item, Descrição, Qtd
        for i, cell in enumerate(tabela.columns[0].cells):
            cell.width = larguras[0]

        cabecalhos_tab = ["Nº", "ITEM", "DESCRIÇÃO E ESPECIFICAÇÃO TÉCNICA", "QTDE"]
        for i, (cell, cab_txt) in enumerate(zip(tabela.rows[0].cells, cabecalhos_tab)):
            cell.text = cab_txt
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Fundo cinza no cabeçalho
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')
            tcPr.append(shd)

        for idx, item in enumerate(itens, 1):
            row = tabela.add_row()
            dados_row = [
                str(idx),
                item.get("nome", ""),
                item.get("descricao", ""),
                str(item.get("quantidade", "")),
            ]
            for i, (cell, texto) in enumerate(zip(row.cells, dados_row)):
                cell.text = texto
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(texto)
                set_font(run, size=10)
                if i in (0, 3):
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        corpo("Não informado.")

    doc.add_paragraph()

    # ── 5. Levantamento de Mercado ───────────────────
    titulo_secao("LEVANTAMENTO DE MERCADO", "5 —")
    corpo(
        "Foram consideradas contratações similares realizadas por outros órgãos e entidades, "
        "com objetivo de identificar soluções que melhor atendam às necessidades da administração, "
        "em conformidade com o art. 18, §1º, VI da Lei 14.133/2021."
    )

    if levantamento:
        for lev in levantamento:
            termo = lev.get("termo", "")
            refs  = lev.get("referencias", [])
            buscado = lev.get("buscado_em", "")

            p = doc.add_paragraph()
            run = p.add_run(f"5.{levantamento.index(lev)+1}. {termo.upper()}")
            set_font(run, size=11, bold=True)
            p.paragraph_format.space_before = Pt(8)

            if refs:
                tab = doc.add_table(rows=1, cols=5)
                tab.style = "Table Grid"

                cabs = ["Órgão", "Município/UF", "Data", "Descrição", "Valor Unit."]
                for cell, cab_txt in zip(tab.rows[0].cells, cabs):
                    cell.text = cab_txt
                    run = cell.paragraphs[0].runs[0]
                    set_font(run, size=9, bold=True)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'D9D9D9')
                    tcPr.append(shd)

                valores = []
                for ref in refs:
                    row = tab.add_row()
                    municipio_ref = ref.get("municipio", "")
                    uf_ref = ref.get("uf", "")
                    loc = f"{municipio_ref}/{uf_ref}" if uf_ref else municipio_ref

                    valor_unit = ref.get("valor_unitario") or ref.get("valor", "")
                    dados_row = [
                        ref.get("orgao", ""),
                        loc,
                        ref.get("data_publicacao", "")[:7] if ref.get("data_publicacao") else "",
                        (ref.get("objeto", "") or ref.get("descricao_item", ""))[:80],
                        str(valor_unit) if valor_unit else "—",
                    ]
                    for i, (cell, texto) in enumerate(zip(row.cells, dados_row)):
                        cell.text = texto
                        r = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(texto)
                        set_font(r, size=9)

                    # Tenta extrair valor numérico para média
                    try:
                        v = str(valor_unit).replace("R$","").replace(".","").replace(",",".").strip()
                        valores.append(float(v))
                    except Exception:
                        pass

                # Linha de estatísticas
                if valores:
                    media = sum(valores) / len(valores)
                    minimo = min(valores)
                    maximo = max(valores)

                    def fmt_brl(v):
                        return f"R$ {v:,.2f}".replace(",","X").replace(".",",").replace("X",".")

                    row_stat = tab.add_row()
                    row_stat.cells[0].text = "ESTATÍSTICAS"
                    run_s = row_stat.cells[0].paragraphs[0].runs[0]
                    set_font(run_s, size=9, bold=True)
                    row_stat.cells[0].merge(row_stat.cells[1])
                    row_stat.cells[2].text = f"Média: {fmt_brl(media)}"
                    row_stat.cells[3].text = f"Mín: {fmt_brl(minimo)}"
                    row_stat.cells[4].text = f"Máx: {fmt_brl(maximo)}"
                    for c in row_stat.cells:
                        r = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else None
                        if r:
                            set_font(r, size=9, bold=True)

                doc.add_paragraph()
    else:
        corpo("Levantamento de mercado não realizado.")

    # ── 6. Estimativa do Valor ───────────────────────
    titulo_secao("ESTIMATIVA DO PREÇO DA CONTRATAÇÃO", "6 —")
    corpo(fixo.get("estimativa_valor", ""))

    # ── 7. Descrição da Solução ──────────────────────
    titulo_secao("DESCRIÇÃO DA SOLUÇÃO COMO UM TODO", "7 —")
    corpo(fixo.get("descricao_solucao", ""))

    # ── 8. Justificativa para Parcelamento ───────────
    titulo_secao("JUSTIFICATIVA PARA O PARCELAMENTO OU NÃO", "8 —")
    corpo(fixo.get("justificativa_parcelamento", ""))

    # ── 9. Resultados Pretendidos ────────────────────
    titulo_secao("DEMONSTRATIVO DOS RESULTADOS PRETENDIDOS", "9 —")
    corpo(fixo.get("resultados_pretendidos", ""))

    # ── 10. Providências Prévias ─────────────────────
    titulo_secao("PROVIDÊNCIAS PRÉVIAS AO CONTRATO", "10 —")
    corpo(fixo.get("providencias_previas", ""))

    # ── 11. Contratações Correlatas ──────────────────
    titulo_secao("CONTRATAÇÕES CORRELATAS/INTERDEPENDENTES", "11 —")
    corpo(fixo.get("contratacoes_correlatas", ""))

    # ── 12. Impactos Ambientais ──────────────────────
    titulo_secao("IMPACTOS AMBIENTAIS", "12 —")
    corpo(fixo.get("impactos_ambientais", ""))

    # ── 13. Viabilidade da Contratação ──────────────
    titulo_secao("POSICIONAMENTO CONCLUSIVO SOBRE A VIABILIDADE", "13 —")
    corpo(fixo.get("posicionamento_conclusivo", ""))

    # ── Campos Livres ────────────────────────────────
    for i, campo in enumerate(campos_livres, 14):
        titulo_secao(campo.get("nome", f"Campo {i}").upper(), f"{i} —")
        corpo(campo.get("conteudo", ""))

    # ── Assinatura ───────────────────────────────────
    doc.add_page_break()
    titulo_secao("APROVAÇÃO E ASSINATURA")

    num_secao = 13 + len(campos_livres) + 1
    corpo(
        f"Com base no exposto acima, a equipe de planejamento declara viável esta contratação, "
        f"salientando tratar-se de solução necessária à continuidade dos serviços prestados pela administração."
    )

    doc.add_paragraph()

    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_local = p_local.add_run(f"{municipio}, {data_fmt}." if municipio else data_fmt + ".")
    set_font(run_local, size=12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Linha de assinatura
    p_linha = doc.add_paragraph()
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linha = p_linha.add_run("_" * 50)
    set_font(run_linha, size=12)

    p_nome = doc.add_paragraph()
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nome = p_nome.add_run(responsavel or "Responsável")
    set_font(run_nome, size=12, bold=True)

    if cargo:
        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cargo = p_cargo.add_run(cargo)
        set_font(run_cargo, size=11)

    # ── Salva em memória ─────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
