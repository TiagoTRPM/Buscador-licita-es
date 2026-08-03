import io
import os
import re
import unicodedata
import tempfile
import zipfile
from lxml import etree

from docx import Document


def pdf_para_docx(pdf_bytes: bytes) -> bytes:
    """Converte arquivo PDF para DOCX usando a biblioteca pdf2docx."""
    from pdf2docx import Converter
    
    # Criar um arquivo temporário com 'delete=False'
    temp_pdf_path = tempfile.mktemp(suffix=".pdf")
    with open(temp_pdf_path, "wb") as f:
        f.write(pdf_bytes)
        
    temp_docx_path = temp_pdf_path + ".docx"
    
    try:
        cv = Converter(temp_pdf_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        
        with open(temp_docx_path, "rb") as f:
            docx_bytes = f.read()
            
        return docx_bytes
    finally:
        if os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception:
                pass


def doc_para_docx(doc_bytes: bytes) -> bytes:
    """Converte arquivo .doc para .docx usando automação COM (win32com) no Windows."""
    import win32com.client
    import pythoncom
    
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_doc:
        temp_doc.write(doc_bytes)
        temp_doc_path = temp_doc.name
        
    temp_doc_path_abs = os.path.abspath(temp_doc_path)
    temp_docx_path_abs = temp_doc_path_abs + "x"
    
    word = None
    doc_obj = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        doc_obj = word.Documents.Open(temp_doc_path_abs)
        doc_obj.SaveAs2(temp_docx_path_abs, FileFormat=16) # 16 = wdFormatXMLDocument
        doc_obj.Close()
        doc_obj = None
        
        with open(temp_docx_path_abs, "rb") as f:
            docx_bytes = f.read()
            
        return docx_bytes
    except Exception as e:
        raise RuntimeError(f"Falha ao converter .doc para .docx. Certifique-se de que o Microsoft Word está instalado. Erro: {e}")
    finally:
        if doc_obj is not None:
            try:
                doc_obj.Close()
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        
        if os.path.exists(temp_doc_path_abs):
            try:
                os.remove(temp_doc_path_abs)
            except Exception:
                pass
        if os.path.exists(temp_docx_path_abs):
            try:
                os.remove(temp_docx_path_abs)
            except Exception:
                pass


def odt_para_docx(odt_bytes: bytes) -> bytes:
    """Converte arquivo ODT para DOCX lendo o XML e gerando parágrafos estruturados."""
    try:
        with zipfile.ZipFile(io.BytesIO(odt_bytes)) as z:
            content_xml = z.read("content.xml")
    except Exception as e:
        raise ValueError(f"Arquivo ODT inválido ou corrompido: {e}")
        
    root = etree.fromstring(content_xml)
    namespaces = {
        'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
        'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
    }
    
    doc_novo = Document()
    elementos = root.xpath('//text:p | //text:h', namespaces=namespaces)
    
    for elem in elementos:
        texto = "".join(elem.itertext()).strip()
        if not texto:
            continue
            
        tag_local = etree.QName(elem).localname
        if tag_local == 'h':
            doc_novo.add_heading(texto, level=1)
        else:
            doc_novo.add_paragraph(texto)
            
    buffer = io.BytesIO()
    doc_novo.save(buffer)
    return buffer.getvalue()


def garantir_docx(filename: str, conteudo: bytes) -> bytes:
    """Garante que o conteúdo fornecido seja retornado em bytes no formato .docx.
    Caso o arquivo seja .pdf, .doc ou .odt, converte-o de forma transparente."""
    ext = (filename or "").lower().split(".")[-1]
    
    if ext == "docx":
        return conteudo
    elif ext == "pdf":
        return pdf_para_docx(conteudo)
    elif ext == "doc":
        return doc_para_docx(conteudo)
    elif ext == "odt":
        return odt_para_docx(conteudo)
    else:
        raise ValueError(f"Formato de arquivo não suportado para conversão: {ext}")


CAMPOS_POR_SECAO = {
    1: "descricao_necessidade",
    2: "previsao_pca",
    3: "requisitos_contratacao",
    6: "estimativa_valor",
    7: "descricao_solucao",
    8: "justificativa_parcelamento",
    9: "resultados_pretendidos",
    10: "providencias_previas",
    11: "contratacoes_correlatas",
    12: "impactos_ambientais",
    13: "posicionamento_conclusivo",
}


def _normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFD", texto or "")
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", texto).strip().lower()


def _numero_secao(texto: str) -> int | None:
    encontrado = re.match(r"^\s*(\d{1,2})\s*(?:[.º°]|[-–—])", texto or "")
    return int(encontrado.group(1)) if encontrado else None


def _eh_titulo(paragrafo) -> bool:
    texto = (paragrafo.text or "").strip()
    if not texto:
        return False
    if _numero_secao(texto) is not None:
        return True

    estilo = _normalizar(getattr(paragrafo.style, "name", ""))
    if "heading" in estilo or "titulo" in estilo or "cabecalho" in estilo:
        return True

    letras = [c for c in texto if c.isalpha()]
    return bool(letras) and len(texto) <= 120 and "".join(letras).isupper()


def extrair_campos_etp(arquivo: bytes) -> dict[str, str]:
    """Extrai os blocos textuais de um ETP Word por numeração e posição."""
    documento = Document(io.BytesIO(arquivo))
    secoes = []
    titulo_atual = None
    conteudo_atual = []

    def finalizar_secao():
        if titulo_atual and conteudo_atual:
            conteudo = "\n\n".join(conteudo_atual).strip()
            if conteudo:
                secoes.append((titulo_atual, conteudo))

    for paragrafo in documento.paragraphs:
        texto = (paragrafo.text or "").strip()
        if not texto:
            continue
        if _eh_titulo(paragrafo):
            finalizar_secao()
            titulo_atual = texto
            conteudo_atual = []
        elif titulo_atual:
            conteudo_atual.append(texto)
    finalizar_secao()

    campos = {}
    usados = set()

    # A numeração é a fonte mais confiável: o nome da seção pode variar.
    for indice, (titulo, conteudo) in enumerate(secoes):
        numero = _numero_secao(titulo)
        campo = CAMPOS_POR_SECAO.get(numero)
        if campo and campo not in campos:
            campos[campo] = conteudo
            usados.add(indice)

    # Para documentos sem numeração, usa a sequência dos blocos de conteúdo.
    candidatos = [
        (indice, conteudo)
        for indice, (titulo, conteudo) in enumerate(secoes)
        if indice not in usados
        and not _normalizar(titulo).startswith(("introducao", "estudo tecnico preliminar"))
    ]
    faltantes = [campo for campo in CAMPOS_POR_SECAO.values() if campo not in campos]
    for campo, (_, conteudo) in zip(faltantes, candidatos):
        campos[campo] = conteudo

    return campos
